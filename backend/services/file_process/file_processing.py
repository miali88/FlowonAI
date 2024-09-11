import logging
from fastapi import UploadFile
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
import io

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

async def process_file(file: UploadFile) -> str:
    logger.info(f"Processing file: {file.filename}")
    
    content = ""
    file_extension = file.filename.split('.')[-1].lower()
    logger.info(f"File extension: {file_extension}")

    try:
        if file_extension == 'pdf':
            print('file is pdf')
            content = await process_pdf(file)
        elif file_extension in ['docx', 'doc']:
            print('file is docx')
            content = await process_docx(file)
        elif file_extension in ['xlsx', 'xls']:
            print('file is excel')
            content = await process_excel(file)
        elif file_extension == 'txt':
            print('file is txt')
            content = (await file.read()).decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"File processed successfully. Content length: {len(content)}")
        return content
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        raise

async def process_pdf(file: UploadFile) -> str:
    logger.info("Processing PDF file")

    pdf_content = await file.read()
    pdf_reader = PdfReader(io.BytesIO(pdf_content))
    content = "\n".join(page.extract_text() for page in pdf_reader.pages)
    logger.info(f"PDF processed. Number of pages: {len(pdf_reader.pages)}")

    print(f"Processed content length: {len(content)} characters")
    print("First 500 characters of content:")
    print(content)

    return content

async def process_docx(file: UploadFile) -> str:
    logger.info("Processing DOCX file")
    docx_content = await file.read()
    doc = Document(io.BytesIO(docx_content))
    content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    logger.info(f"DOCX processed. Number of paragraphs: {len(doc.paragraphs)}")
    return content

async def process_excel(file: UploadFile) -> str:
    logger.info("Processing Excel file")
    excel_content = await file.read()
    workbook = load_workbook(filename=io.BytesIO(excel_content))
    content = []
    for sheet in workbook.sheetnames:
        ws = workbook[sheet]
        content.append(f"Sheet: {sheet}")
        for row in ws.iter_rows(values_only=True):
            content.append("\t".join(str(cell) for cell in row if cell is not None))
    logger.info(f"Excel processed. Number of sheets: {len(workbook.sheetnames)}")
    return "\n".join(content)