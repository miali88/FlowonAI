import logging
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook # type: ignore
import io
from typing import Any
import csv  # Add this import at the top

from fastapi import UploadFile, BackgroundTasks

from services.supabase.client import get_supabase
from services.knowledge_base.vectorise_data import kb_item_to_chunks

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

async def process_file(file: UploadFile) -> Any:
    logger.info(f"Processing file: {file.filename}")

    content = ""
    if not file.filename:
        raise ValueError("File name is required")
    file_extension = file.filename.split('.')[-1].lower()
    logger.info(f"File extension: {file_extension}")

    try:
        if file_extension == 'pdf':
            logger.info('Processing PDF file type')
            content = await process_pdf(file)
            return content, file_extension, False  # False indicates non-tabular data
        elif file_extension in ['docx', 'doc']:
            logger.info('Processing DOCX file type')
            content = await process_docx(file)
        elif file_extension in ['xlsx', 'xls']:
            logger.info('Processing Excel file type')
            content = await process_excel(file)
            return content, file_extension, True   # True indicates tabular data
        elif file_extension == 'txt':
            logger.info('Processing TXT file type')
            content = (await file.read()).decode('utf-8')
        elif file_extension == 'csv':
            logger.info('Processing CSV file type')
            content = await process_csv(file)
            return content, file_extension, True
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        logger.info(f"File processed successfully. Content length: {len(content)}")
        return content, file_extension, False  # Default to non-tabular
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        raise

async def process_pdf(file: UploadFile) -> str:
    logger.info("Processing PDF file")

    pdf_content = await file.read()
    pdf_reader = PdfReader(io.BytesIO(pdf_content))
    content = "\n".join(page.extract_text() for page in pdf_reader.pages)
    logger.info(f"PDF processed. Number of pages: {len(pdf_reader.pages)}")
    return content

async def process_docx(file: UploadFile) -> str:
    logger.info("Processing DOCX file")
    docx_content = await file.read()
    doc = Document(io.BytesIO(docx_content))
    content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    logger.info(f"DOCX processed. Number of paragraphs: {len(doc.paragraphs)}")
    return content

async def process_excel(file: UploadFile) -> list[dict]:
    logger.info("Processing Excel file")
    excel_content = await file.read()
    workbook = load_workbook(filename=io.BytesIO(excel_content))
    rows = []
    
    for sheet in workbook.sheetnames:
        ws = workbook[sheet]
        # Get headers from first row
        headers = [str(cell.value) if cell.value is not None else f"column_{idx}" 
                  for idx, cell in enumerate(next(ws.rows))]
        
        # Process remaining rows
        for row in list(ws.iter_rows(min_row=2, values_only=True)):
            row_dict = {
                "sheet_name": sheet,
                **{headers[i]: str(val) if val is not None else "" 
                   for i, val in enumerate(row)}
            }
            rows.append(row_dict)
            
    logger.info(f"Excel processed. Total rows: {len(rows)}")
    return rows

async def process_csv(file: UploadFile) -> list[dict]:
    logger.info("Processing CSV file")
    csv_content = await file.read()
    
    # Decode bytes to string and create a string IO object
    csv_string = csv_content.decode('utf-8')
    csv_file = io.StringIO(csv_string)
    
    # Read CSV file into a list of dictionaries
    csv_reader = csv.DictReader(csv_file)
    content = [row for row in csv_reader]  # Each row is now a dictionary

    logger.info("CSV processed successfully")
    return content  # Returns a list of JSON-like dictionaries

""" ENTRY POINT """
async def process_and_store_file(
    file: UploadFile,
    user_id: str,
    background_tasks: BackgroundTasks
) -> dict:
    """Process file and store in database with background chunking"""
    logger.info(f"Processing and storing file: {file.filename}")
    
    # Process the file
    content, file_extension, is_tabular = await process_file(file)
    
    # Get supabase client
    supabase = await get_supabase()
    
    """ Vectorising data """
    if is_tabular:
        # Create parent record
        parent_item = await supabase.table('user_text_files').insert({
            "title": file.filename,
            "heading": file.filename,
            "file_name": file.filename,
            "content": "",  # Empty content since data is in child rows
            "user_id": user_id,
            "data_type": file_extension
        }).execute()
        
        parent_id = parent_item.data[0]['id']
        
        # Batch insert rows into a new table for tabular data
        rows_to_insert = [{
            "file_id": parent_id,
            "user_id": user_id,
            "row_data": row,
            "file_name": file.filename,
        } for row in content]
        
        # Insert in batches of 1000 to avoid hitting limits
        BATCH_SIZE = 1000
        for i in range(0, len(rows_to_insert), BATCH_SIZE):
            batch = rows_to_insert[i:i + BATCH_SIZE]
            await supabase.table('user_tabular_data').insert(batch).execute()
        
        # Schedule chunking task
        background_tasks.add_task(
            kb_item_to_chunks,
            parent_id,  # for tabular data
            rows_to_insert,  # pass the rows directly
            user_id,
            file.filename,
            is_tabular=True
        )
        
        return parent_item.data[0]
    else:
        # Insert into user_text_files
        new_item = await supabase.table('user_text_files').insert({
            "title": file.filename,
            "heading": file.filename,
            "file_name": file.filename,
            "content": content,
            "user_id": user_id,
            "data_type": file_extension
        }).execute()
        
        # Insert into headers table
        await supabase.table('user_text_files_headers').insert({
            "heading": file.filename,
            "file_name": file.filename,
            "user_id": user_id,
            "data_type": file_extension,
            "parent_id": new_item.data[0]['id']
        }).execute()
        
        # Schedule chunking task
        background_tasks.add_task(
            kb_item_to_chunks,
            new_item.data[0]['id'],
            content,
            user_id,
            new_item.data[0]['title']
        )
        
        return new_item.data[0]
